from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable
from uuid import uuid4

import pdfplumber
import pymupdf
from docx import Document as DocxDocument
from openpyxl import load_workbook
from PIL import Image
import pytesseract

from app.services.ollama_client import caption_image


@dataclass
class ParsedBlock:
    text: str
    page: int | None = None
    sheet: str | None = None
    kind: str = "text"


async def parse_file(path: str, content_type: str, filename: str) -> list[ParsedBlock]:
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix == ".pdf" or "pdf" in content_type:
        return await _parse_pdf(p)
    if suffix in {".docx", ".doc"} or "word" in content_type:
        return _parse_docx(p)
    if suffix in {".xlsx", ".xls"} or "sheet" in content_type or "excel" in content_type:
        return _parse_xlsx(p)
    if suffix in {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff"} or content_type.startswith("image/"):
        return await _parse_image(p)
    raise ValueError(f"Unsupported file type: {filename}")


async def _parse_pdf(path: Path) -> list[ParsedBlock]:
    blocks: list[ParsedBlock] = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                blocks.append(ParsedBlock(text=text, page=i, kind="text"))
            for table in page.extract_tables() or []:
                rows = [" | ".join((c or "").strip() for c in row) for row in table if row]
                if rows:
                    md = "Table on page {p}:\n".format(p=i) + "\n".join(rows)
                    blocks.append(ParsedBlock(text=md, page=i, kind="table"))
    doc = pymupdf.open(path)
    for i, page in enumerate(doc, start=1):
        for img in page.get_images(full=True):
            xref = img[0]
            try:
                pix = pymupdf.Pixmap(doc, xref)
                if pix.n - pix.alpha > 3:
                    pix = pymupdf.Pixmap(pymupdf.csRGB, pix)
                img_bytes = pix.tobytes("png")
            except Exception:
                continue
            ocr = ""
            try:
                from io import BytesIO

                im = Image.open(BytesIO(img_bytes))
                ocr = pytesseract.image_to_string(im) or ""
            except Exception:
                ocr = ""
            caption = await caption_image(img_bytes)
            parts = [p for p in (caption.strip(), ocr.strip()) if p]
            if parts:
                blocks.append(
                    ParsedBlock(
                        text="Figure on page {p}:\n".format(p=i) + "\n".join(parts),
                        page=i,
                        kind="figure",
                    )
                )
    return blocks or [ParsedBlock(text="(empty PDF)", page=1)]


def _parse_docx(path: Path) -> list[ParsedBlock]:
    doc = DocxDocument(str(path))
    blocks: list[ParsedBlock] = []
    buffer: list[str] = []
    for para in doc.paragraphs:
        style = (para.style.name or "") if para.style else ""
        text = para.text.strip()
        if not text:
            continue
        if "Heading" in style and buffer:
            blocks.append(ParsedBlock(text="\n".join(buffer), kind="text"))
            buffer = [text]
        else:
            buffer.append(text)
    if buffer:
        blocks.append(ParsedBlock(text="\n".join(buffer), kind="text"))
    for table in doc.tables:
        rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        if rows:
            blocks.append(ParsedBlock(text="Table:\n" + "\n".join(rows), kind="table"))
    return blocks or [ParsedBlock(text="(empty document)")]


def _parse_xlsx(path: Path) -> list[ParsedBlock]:
    wb = load_workbook(path, data_only=True, read_only=True)
    blocks: list[ParsedBlock] = []
    for sheet in wb.worksheets:
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue
        headers = [str(c).strip() if c is not None else "" for c in rows[0]]
        header_line = " | ".join(headers)
        group: list[str] = []
        for row in rows[1:]:
            vals = [str(c).strip() if c is not None else "" for c in row]
            if not any(vals):
                continue
            paired = ", ".join(f"{h}: {v}" for h, v in zip(headers, vals) if h or v)
            group.append(paired)
            if len(group) >= 12:
                blocks.append(
                    ParsedBlock(
                        text=f"Sheet {sheet.title}\nColumns: {header_line}\n" + "\n".join(group),
                        sheet=sheet.title,
                        kind="table",
                    )
                )
                group = []
        if group:
            blocks.append(
                ParsedBlock(
                    text=f"Sheet {sheet.title}\nColumns: {header_line}\n" + "\n".join(group),
                    sheet=sheet.title,
                    kind="table",
                )
            )
    return blocks or [ParsedBlock(text="(empty spreadsheet)", kind="table")]


async def _parse_image(path: Path) -> list[ParsedBlock]:
    data = path.read_bytes()
    ocr = ""
    try:
        ocr = pytesseract.image_to_string(Image.open(path)) or ""
    except Exception:
        ocr = ""
    caption = await caption_image(data)
    text = "\n".join(p for p in (caption.strip(), ocr.strip()) if p) or "(image with no extracted text)"
    return [ParsedBlock(text=text, kind="figure")]
